"""
Format-aware loaders. Each returns RawUnits — the smallest structural piece
worth treating separately (a headed section, a table row, an OAS operation,
a spec summary, an inventory entry).

OCR: images embedded in .docx are read via an OpenAI-compatible vision
call (OCR_MODEL if set, otherwise CHAT_MODEL) — not a local OCR binary.
The extracted text is APPENDED to the section the image sits in (an image
illustrates the rule around it, so it must not float as its own chunk). If
the vision call fails (model doesn't support images, network error, etc.),
ingestion still works and just skips that image's text with a warning.
"""
import base64
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import yaml
from docx import Document as DocxDocument
from pypdf import PdfReader

logger = logging.getLogger(__name__)


@dataclass
class RawUnit:
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)


_OCR_PROMPT = (
    "Extract all readable text from this image verbatim, preserving line "
    "breaks. If the image is a diagram, also briefly describe its structure "
    "(boxes/arrows/flow) in one sentence before the extracted text. Return "
    "only that — no commentary, no markdown fences."
)


def _ocr_image(image_bytes: bytes, content_type: str = "image/png") -> str:
    """Read embedded-image text via a vision-capable OpenAI-compatible call
    — not a local OCR binary. Uses OCR_MODEL if set (a dedicated internal
    OCR deployment), otherwise falls back to CHAT_MODEL."""
    try:
        from langchain_core.messages import HumanMessage
        from app.integrations.internal_llm import get_ocr_model

        b64 = base64.b64encode(image_bytes).decode()
        message = HumanMessage(content=[
            {"type": "text", "text": _OCR_PROMPT},
            {"type": "image_url", "image_url": {"url": f"data:{content_type};base64,{b64}"}},
        ])
        response = get_ocr_model().invoke([message])
        return (response.content or "").strip()
    except Exception:
        logger.warning("Vision OCR call failed; skipping this image's text", exc_info=True)
        return ""


# ---------------------------------------------------------------- docx ----

def load_docx(path: str | Path) -> list[RawUnit]:
    """
    Heading-based sections (prose), flattened table rows, and OCR'd images
    appended to the section they appear in.
    """
    doc = DocxDocument(str(path))
    source = Path(path).name
    units: list[RawUnit] = []

    # Map image relationship id -> (bytes, content_type), so inline images
    # can be sent to the vision model with the right data URI mime type.
    image_parts = {
        rel_id: (rel.target_part.blob, rel.target_part.content_type)
        for rel_id, rel in doc.part.rels.items()
        if "image" in rel.reltype
    }

    current_heading = "Introduction"
    current_text: list[str] = []
    current_has_ocr = False

    def flush():
        nonlocal current_has_ocr
        body = "\n".join(current_text).strip()
        if body:
            units.append(RawUnit(
                text=body,
                metadata={"source": source, "section": current_heading,
                          "type": "prose", "has_ocr": current_has_ocr},
            ))
        current_has_ocr = False

    for para in doc.paragraphs:
        style_name = para.style.name if para.style is not None else ""
        if style_name.startswith("Heading"):
            flush()
            current_heading = para.text.strip() or current_heading
            current_text = []
            continue

        if para.text.strip():
            current_text.append(para.text.strip())

        # OCR any images anchored in this paragraph; append text to the section.
        for run in para.runs:
            for blip in run._element.findall(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            ):
                rel_id = blip.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                )
                if rel_id in image_parts:
                    image_bytes, content_type = image_parts[rel_id]
                    ocr_text = _ocr_image(image_bytes, content_type)
                    if ocr_text:
                        current_text.append(f"[image text] {ocr_text}")
                        current_has_ocr = True
    flush()

    # Tables: one flattened row per unit ("Header1: value. Header2: value.")
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [c.text.strip() for c in table.rows[0].cells]
        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            sentence = ". ".join(f"{h}: {v}" for h, v in zip(headers, cells) if v)
            if sentence:
                units.append(RawUnit(
                    text=sentence,
                    metadata={"source": source, "type": "table_row"},
                ))
    return units


# ----------------------------------------------------------------- pdf ----

def load_pdf(path: str | Path) -> list[RawUnit]:
    reader = PdfReader(str(path))
    return [
        RawUnit(text=page.extract_text() or "",
                metadata={"source": Path(path).name, "page": i + 1, "type": "prose"})
        for i, page in enumerate(reader.pages)
    ]


# ----------------------------------------------------------------- oas ----

def load_oas(path: str | Path, api_id: str | None = None) -> list[RawUnit]:
    """
    One unit per (path, method) operation + ONE spec-summary unit.
    Every unit carries the same api_id — the link that keeps all chunks
    of one spec (and its referential entry) connected.
    """
    spec = yaml.safe_load(Path(path).read_text())
    info = spec.get("info", {})
    api_name = info.get("title", Path(path).stem)
    api_id = api_id or Path(path).stem  # convention: filename stem unless given

    units: list[RawUnit] = []
    endpoint_list: list[str] = []

    for route, methods in (spec.get("paths") or {}).items():
        for method, op in (methods or {}).items():
            if method not in {"get", "post", "put", "patch", "delete"}:
                continue
            endpoint_list.append(f"{method.upper()} {route}")
            params = op.get("parameters", [])
            param_text = "\n".join(
                f"  - {p.get('name')} ({p.get('in')}, required={p.get('required', False)})"
                for p in params if isinstance(p, dict)
            )
            units.append(RawUnit(
                text=(f"{method.upper()} {route}\n"
                      f"summary: {op.get('summary', '')}\n"
                      f"{op.get('description', '')}\n"
                      f"parameters:\n{param_text}"),
                metadata={"source": Path(path).name, "api_id": api_id,
                          "api_name": api_name, "path": route,
                          "method": method.upper(), "type": "oas_operation"},
            ))

    # Spec summary chunk — the "table of contents" for the whole API.
    units.append(RawUnit(
        text=(f"API: {api_name} (version {info.get('version', '?')})\n"
              f"{info.get('description', '')}\n"
              f"Endpoints:\n" + "\n".join(f"  - {e}" for e in endpoint_list)),
        metadata={"source": Path(path).name, "api_id": api_id,
                  "api_name": api_name, "type": "oas_summary"},
    ))
    return units


# --------------------------------------------------------- referential ----

def load_referential(path: str | Path) -> list[RawUnit]:
    """
    API Referential inventory — expects a YAML/JSON list of entries:
      - api_id: doc-mgmt-api
        name: Document Management API
        description: ...
        url: ...
    One unit per API, no splitting (an entry is already atomic).
    Swap this for a REST call to the live Referential when ready.
    """
    entries = yaml.safe_load(Path(path).read_text()) or []
    units = []
    for e in entries:
        text = f"{e.get('name', '')}: {e.get('description', '')}"
        extras = {k: v for k, v in e.items() if k not in {"name", "description"} and v}
        if extras:
            text += "\n" + "\n".join(f"{k}: {v}" for k, v in extras.items())
        units.append(RawUnit(
            text=text,
            metadata={"source": Path(path).name, "api_id": e.get("api_id", ""),
                      "api_name": e.get("name", ""), "url": e.get("url"),
                      "type": "referential_entry"},
        ))
    return units
